"""
document_service.py — Document Business Logic
==============================================
Handles:
  - Document CRUD (create, list, get, update, delete)
  - File upload (validate extension, extract text from PDF/DOCX/TXT)
  - Text chunking (split into ~800 char chunks with overlap)
  - Embedding generation
  - Document reprocessing
  - Category and Status management
  - AI-generated document summary

RULES:
  - Calls Repository for DB operations
  - Calls ai_agent for embedding and summary generation
  - No direct DB queries
  - No HTTP route concerns
"""

import io
import os
import re
import time
import uuid
from datetime import datetime, timezone
from typing import Optional, List
from uuid import UUID

from fastapi import HTTPException, UploadFile, status
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.core.logger import logger
from app.models.models import (
    Document,
    DocumentVersion,
    DocumentChunk,
    DocumentProcessingLog,
    DocumentCategory,
)
from app.repositories.document_repository import (
    DocumentRepository,
    DocumentVersionRepository,
    DocumentChunkRepository,
    DocumentProcessingLogRepository,
    DocumentCategoryRepository,
    DocumentStatusRepository,
)
from app.schemas.schemas import (
    DocumentCreate,
    DocumentUpdate,
    CategoryCreate,
    CategoryUpdate,
)
from app.ai import ai_agent


# ─────────────────────────────────────────────────────────────────────────────
# TEXT EXTRACTION UTILITIES
# ─────────────────────────────────────────────────────────────────────────────

async def _extract_text_from_file(file: UploadFile):
    """
    Extracts plain text from uploaded PDF, DOCX, or TXT file.

    Returns:
        Tuple: (extracted_text, page_count)
    """
    content = await file.read()
    ext = file.filename.rsplit(".", 1)[-1].lower()

    logger.debug(f"[DocumentService] Extracting text from .{ext} file: {file.filename}")

    if ext not in settings.ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Unsupported file type '.{ext}'. Allowed: {settings.ALLOWED_EXTENSIONS}",
        )

    try:
        if ext == "txt":
            return content.decode("utf-8", errors="replace"), 1

        if ext == "pdf":
            import pypdf

            reader = pypdf.PdfReader(io.BytesIO(content))
            pages = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text.strip())

            return "\n\n".join(pages), len(reader.pages)

        if ext == "docx":
            import docx

            doc = docx.Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

            return "\n\n".join(paragraphs), None

    except Exception as exc:
        logger.error(f"[DocumentService] Text extraction failed: {exc}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Failed to extract text from the file: {str(exc)}",
        )

    return "", None


def _chunk_text(text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
    """
    Splits long text into overlapping chunks.
    """
    chunk_size = chunk_size or settings.CHUNK_SIZE
    overlap = overlap or settings.CHUNK_OVERLAP

    if not text:
        return []

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end].strip()

        if chunk:
            chunks.append(chunk)

        start += chunk_size - overlap

    logger.debug(f"[DocumentService] Chunked text into {len(chunks)} chunks")
    return chunks


async def _save_file_to_disk(file: UploadFile, file_content: bytes) -> str:
    """
    Saves uploaded file to local uploads directory.
    """
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)

    unique_name = f"{uuid.uuid4()}_{file.filename}"
    file_path = os.path.join(settings.UPLOAD_DIR, unique_name)

    with open(file_path, "wb") as f:
        f.write(file_content)

    logger.debug(f"[DocumentService] File saved to: {file_path}")
    return file_path


def _is_missing_summary(summary: Optional[str]) -> bool:
    """
    Checks whether summary is actually missing or fallback text.
    """
    if not summary:
        return True

    cleaned = summary.strip().lower()

    return cleaned in {
        "",
        "no summary available for this document.",
        "no summary available for this document",
        "no summary available",
    }


def _clean_markdown_to_plain_text(text: str) -> str:
    """
    Converts markdown-looking AI response into normal plain text.
    """
    if not text:
        return ""

    text = re.sub(r"```[\s\S]*?```", "", text)
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"$begin:math:display$\(\[\^$end:math:display$]+)\]$begin:math:text$\[\^\)\]\+$end:math:text$", r"\1", text)
    text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*>\s?", "", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


async def _generate_ai_document_summary(summary_input: str) -> Optional[str]:
    """
    Generates long detailed document summary using AI.

    Important:
    - This function tries to call generate_document_summary with log_response=False.
    - If your ai_agent.py does not support log_response yet, add it there.
    """
    if not summary_input or not summary_input.strip():
        return None

    prompt = (
        "Give detailed description of the document.\n\n"
        "Read the following document content and write a long and detailed plain text "
        "description of this document/project. Explain what it is about, its purpose, "
        "important modules, workflow, technologies, architecture, roles, database or API "
        "flow if mentioned, and how it helps the user. Write it for someone opening the "
        "document for the first time. Do not use markdown, headings, bullet points, "
        "tables, bold text, or code blocks. Use normal plain text only. Keep the answer "
        "in 2 to 4 detailed paragraphs.\n\n"
        f"Document content:\n{summary_input}"
    )

    if hasattr(ai_agent, "generate_document_summary"):
        try:
            return await ai_agent.generate_document_summary(
                summary_input,
                log_response=False,
            )
        except TypeError:
            return await ai_agent.generate_document_summary(summary_input)

    if hasattr(ai_agent, "generate_chat_response"):
        return await ai_agent.generate_chat_response(
            system_prompt=(
                "You are an expert document summarizer. Return plain text only. "
                "Do not use markdown, headings, bullet points, tables, bold text, or code blocks."
            ),
            user_prompt=prompt,
        )

    logger.warning("[DocumentService] No AI summary generation function found in ai_agent.py")
    return None


# ─────────────────────────────────────────────────────────────────────────────
# CATEGORY SERVICE
# ─────────────────────────────────────────────────────────────────────────────

class CategoryService:
    """Handles document category CRUD operations."""

    def __init__(self, db: AsyncSession) -> None:
        self._repo = DocumentCategoryRepository(db)

    async def create_category(self, data: CategoryCreate, user_id: UUID) -> DocumentCategory:
        existing = await self._repo.get_by_name(data.name)

        if existing:
            raise HTTPException(
                status_code=status.HTTP_409_CONFLICT,
                detail=f"Category '{data.name}' already exists.",
            )

        cat = DocumentCategory(
            name=data.name,
            description=data.description,
            created_by=user_id,
        )

        return await self._repo.create(cat)

    async def list_categories(self, page, page_size, search, is_active, sort_by, sort_order):
        return await self._repo.list_categories(
            page,
            page_size,
            search,
            is_active,
            sort_by,
            sort_order,
        )

    async def get_category(self, cat_id: UUID) -> DocumentCategory:
        cat = await self._repo.get_by_id(cat_id)

        if not cat:
            raise HTTPException(status_code=404, detail="Category not found.")

        return cat

    async def update_category(
        self,
        cat_id: UUID,
        data: CategoryUpdate,
        user_id: UUID,
    ) -> DocumentCategory:
        cat = await self.get_category(cat_id)

        update_data = {
            key: value
            for key, value in data.model_dump().items()
            if value is not None
        }
        update_data["updated_by"] = user_id

        return await self._repo.update(cat, update_data)

    async def delete_category(self, cat_id: UUID) -> DocumentCategory:
        cat = await self.get_category(cat_id)
        return await self._repo.soft_delete(cat)


# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT SERVICE
# ─────────────────────────────────────────────────────────────────────────────

class DocumentService:
    """
    Core service for document upload, processing, and management.
    Coordinates between repositories and AI agent layer.
    """

    def __init__(self, db: AsyncSession) -> None:
        self._db = db
        self._doc_repo = DocumentRepository(db)
        self._ver_repo = DocumentVersionRepository(db)
        self._chunk_repo = DocumentChunkRepository(db)
        self._log_repo = DocumentProcessingLogRepository(db)
        self._status_repo = DocumentStatusRepository(db)
    # ── Create ────────────────────────────────────────────────────────────────
    async def upload_document(
        self,
        metadata: DocumentCreate,
        file: Optional[UploadFile],
        text_content: Optional[str],
        user_id: UUID,
    ) -> Document:
        logger.info(f"[DocumentService] Uploading document: '{metadata.title}'")

        if not file and not text_content:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Provide either a file upload or paste text content.",
            )
        if file:
            file_bytes = await file.read()
            await file.seek(0)

            size_mb = len(file_bytes) / (1024 * 1024)

            if size_mb > settings.MAX_UPLOAD_SIZE_MB:
                raise HTTPException(
                    status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                    detail=f"File too large. Max size is {settings.MAX_UPLOAD_SIZE_MB} MB.",
                )

        pending_status = await self._status_repo.get_by_name("pending")

        if not pending_status:
            raise HTTPException(
                status_code=500,
                detail="'pending' status not seeded in DB.",
            )

        doc = Document(
            title=metadata.title,
            category_id=metadata.category_id,
            tags=metadata.tags,
            uploaded_by=user_id,
            status_id=pending_status.id,
            current_version_no=1,
            created_by=user_id,
        )

        doc = await self._doc_repo.create(doc)
        extracted_text = text_content
        page_count = None if not text_content else 1
        file_name = None
        file_type = None
        file_path = None

        if file:
            extracted_text, page_count = await _extract_text_from_file(file)

            file_name = file.filename
            file_type = file.filename.rsplit(".", 1)[-1].lower()

            await file.seek(0)
            file_bytes = await file.read()
            file_path = await _save_file_to_disk(file, file_bytes)

        token_count = None

        if extracted_text:
            token_count = max(1, round(len(extracted_text.split()) * 1.3))

        version = DocumentVersion(
            document_id=doc.id,
            version_no=1,
            version_label="v1.0",
            title=metadata.title,
            content_text=extracted_text,
            file_name=file_name,
            file_type=file_type,
            file_path=file_path,
            page_count=page_count,
            token_count=token_count,
            is_current=True,
            status_id=pending_status.id,
            created_by=user_id,
        )

        await self._ver_repo.create(version)

        logger.info(f"[DocumentService] Document created: {doc.id}")
        return await self.get_document(doc.id)

    async def upload_new_version(
        self,
        document_id: UUID,
        title: Optional[str],
        version_label: Optional[str],
        change_note: Optional[str],
        file: Optional[UploadFile],
        text_content: Optional[str],
        user_id: UUID,
    ) -> DocumentVersion:
        doc = await self.get_document(document_id)
        if not file and not text_content:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Provide either a file upload or paste text content for the new version.",
            )

        if file:
            file_bytes = await file.read()
            await file.seek(0)

            size_mb = len(file_bytes) / (1024 * 1024)

            if size_mb > settings.MAX_UPLOAD_SIZE_MB:
                raise HTTPException(
                    status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                    detail=f"File too large. Max size is {settings.MAX_UPLOAD_SIZE_MB} MB.",
                )

        pending_status = await self._status_repo.get_by_name("pending")

        if not pending_status:
            raise HTTPException(
                status_code=500,
                detail="'pending' status not seeded in DB.",
            )

        extracted_text = text_content
        page_count = None if not text_content else 1
        file_name = None
        file_type = None
        file_path = None

        if file:
            extracted_text, page_count = await _extract_text_from_file(file)

            file_name = file.filename
            file_type = file.filename.rsplit(".", 1)[-1].lower()

            await file.seek(0)
            file_bytes = await file.read()
            file_path = await _save_file_to_disk(file, file_bytes)

        token_count = None

        if extracted_text:
            token_count = max(1, round(len(extracted_text.split()) * 1.3))

        next_version_no = (await self._ver_repo.get_max_version_no(document_id)) + 1
        await self._ver_repo.set_all_non_current(document_id)

        version = DocumentVersion(
            document_id=document_id,
            version_no=next_version_no,
            version_label=version_label or f"v{next_version_no}.0",
            title=title or doc.title,
            content_text=extracted_text,
            file_name=file_name,
            file_type=file_type,
            file_path=file_path,
            page_count=page_count,
            token_count=token_count,
            change_note=change_note,
            is_current=True,
            status_id=pending_status.id,
            created_by=user_id,
        )

        version = await self._ver_repo.create(version)

        await self._doc_repo.update(
            doc,
            {
                "title": title or doc.title,
                "status_id": pending_status.id,
                "current_version_no": next_version_no,
                "updated_by": user_id,
            },
        )

        refreshed_version = await self._ver_repo.get_current_version(document_id)
        return refreshed_version or version

    # ── Process ───────────────────────────────────────────────────────────────

    async def process_document(self, document_id: UUID, user_id: UUID) -> dict:
        logger.info(f"[DocumentService] Processing document: {document_id}")

        doc = await self._doc_repo.get_with_relations(document_id)

        if not doc:
            raise HTTPException(status_code=404, detail="Document not found.")

        version = await self._ver_repo.get_current_version(document_id)

        if not version or not version.content_text:
            raise HTTPException(
                status_code=422,
                detail="Document version has no text content to process.",
            )

        processing_status = await self._status_repo.get_by_name("processing")

        if not processing_status:
            raise HTTPException(
                status_code=500,
                detail="'processing' status not seeded in DB.",
            )

        proc_log = DocumentProcessingLog(
            document_version_id=version.id,
            status="queued",
            created_by=user_id,
        )

        proc_log = await self._log_repo.create(proc_log)

        start_time = time.time()

        try:
            await self._log_repo.update(proc_log, {"status": "processing"})

            await self._doc_repo.update(
                doc,
                {
                    "status_id": processing_status.id,
                    "updated_by": user_id,
                },
            )

            await self._ver_repo.update(
                version,
                {
                    "status_id": processing_status.id,
                    "updated_by": user_id,
                },
            )

            old_chunks = await self._chunk_repo.get_chunks_by_version(version.id)
            if old_chunks:
                await self._chunk_repo.delete_chunks_by_version(version.id)

            text_chunks = _chunk_text(version.content_text)

            if not text_chunks:
                raise ValueError("No text extracted after chunking.")

            embeddings = await ai_agent.generate_embeddings(text_chunks)

            db_chunks = []

            for index, chunk_text in enumerate(text_chunks, start=1):
                db_chunks.append(
                    DocumentChunk(
                        document_version_id=version.id,
                        chunk_no=index,
                        chunk_text=chunk_text,
                        keyword_text=" ".join(chunk_text.lower().split()[:20]),
                        embedding=embeddings[index - 1],
                        created_by=user_id,
                    )
                )

            await self._chunk_repo.create_many(db_chunks)

            processed_status = await self._status_repo.get_by_name("processed")

            if not processed_status:
                raise HTTPException(
                    status_code=500,
                    detail="'processed' status not seeded in DB.",
                )

            now = datetime.now(timezone.utc)

            await self._doc_repo.update(
                doc,
                {
                    "status_id": processed_status.id,
                    "updated_by": user_id,
                },
            )

            await self._ver_repo.update(
                version,
                {
                    "status_id": processed_status.id,
                    "updated_by": user_id,
                },
            )
            # Generate and store AI summary during processing.
            await self._generate_and_store_summary(
                document=doc,
                version=version,
                chunks=text_chunks,
            )

            await self._log_repo.update(
                proc_log,
                {
                    "status": "success",
                    "processed_at": now,
                },
            )

            elapsed = round((time.time() - start_time) * 1000)

            logger.info(
                f"[DocumentService] Document {document_id} processed: "
                f"{len(db_chunks)} chunks in {elapsed}ms"
            )

            return {
                "document_id": str(document_id),
                "version_no": version.version_no,
                "chunk_count": len(db_chunks),
                "time_taken_ms": elapsed,
            }

        except Exception as exc:
            logger.error(
                f"[DocumentService] Processing failed for {document_id}: {exc}",
                exc_info=True,
            )

            failed_status = await self._status_repo.get_by_name("failed")

            if failed_status:
                await self._doc_repo.update(
                    doc,
                    {
                        "status_id": failed_status.id,
                        "updated_by": user_id,
                    },
                )

            await self._log_repo.update(
                proc_log,
                {
                    "status": "failed",
                    "error_message": str(exc),
                    "processed_at": datetime.now(timezone.utc),
                },
            )

            raise HTTPException(
                status_code=500,
                detail=f"Document processing failed: {str(exc)}",
            )

    async def _generate_and_store_summary(
        self,
        document: Document,
        version: DocumentVersion,
        chunks: List[str],
    ) -> str:
        """
        Generates AI summary and stores it in current document version.
        Used during processing.
        """
        try:
            summary_input = "\n\n".join(chunks[:20])[:12000]

            summary = await _generate_ai_document_summary(summary_input)

            if _is_missing_summary(summary):
                summary = (
                    getattr(document, "description", None)
                    or summary_input[:1000]
                    or "No summary available for this document."
                )
            summary = _clean_markdown_to_plain_text(summary)
            await self._ver_repo.update(
                version,
                {
                    "summary": summary,
                },
            )

            return summary

        except Exception as exc:
            logger.warning(
                f"[DocumentService] Failed to generate automatic summary "
                f"for {document.id}: {exc}"
            )

            fallback_summary = (
                getattr(document, "description", None)
                or getattr(version, "content_text", "")[:1000]
                or "No summary available for this document."
            )

            fallback_summary = _clean_markdown_to_plain_text(fallback_summary)

            await self._ver_repo.update(
                version,
                {
                    "summary": fallback_summary,
                },
            )

            return fallback_summary

    # ── Publish / QA ──────────────────────────────────────────────────────────

    async def publish_document(self, document_id: UUID, user_id: UUID) -> Document:
        doc = await self.get_document(document_id)

        version = await self._ver_repo.get_current_version(document_id)

        if not version:
            raise HTTPException(
                status_code=404,
                detail="Current document version not found.",
            )

        processed_status = await self._status_repo.get_by_name("processed")
        published_status = await self._status_repo.get_by_name("published")

        if not processed_status or not published_status:
            raise HTTPException(
                status_code=500,
                detail="Document statuses are not seeded correctly.",
            )

        if version.status_id != processed_status.id and doc.status_id != processed_status.id:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Only processed documents can be published. Process the document first.",
            )

        if version.qa_test_status == "not_tested":
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Ask a test question before publishing this document.",
            )

        if version.qa_test_status == "needs_fix":
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="This document needs fixes before publishing.",
            )

        if version.qa_test_status != "looks_good":
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Mark the test answer as Looks Good before publishing.",
            )

        chunks = await self._chunk_repo.get_chunks_by_version(version.id)

        if not chunks:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Document has no chunks. Processing might have failed or the document is empty.",
            )

        await self._ver_repo.update(
            version,
            {
                "status_id": published_status.id,
                "updated_by": user_id,
            },
        )

        await self._doc_repo.update(
            doc,
            {
                "status_id": published_status.id,
                "updated_by": user_id,
            },
        )

        return await self.get_document(document_id)

    async def update_qa_status(
        self,document_id: UUID,status_value: str,
        user_id: UUID,
        qa_question: Optional[str] = None,
        qa_answer: Optional[str] = None,
        qa_sources: Optional[List[dict]] = None,
    ) -> Document:
        doc = await self.get_document(document_id)
        version = await self._ver_repo.get_current_version(document_id)

        if not version:
            raise HTTPException(
                status_code=404,
                detail="Current document version not found.",
            )

        if status_value not in ["looks_good", "needs_fix"]:
            raise HTTPException(
                status_code=400,
                detail="Invalid status. Must be 'looks_good' or 'needs_fix'",
            )

        import json

        update_data = {
            "qa_test_status": status_value,
            "qa_tested_at": datetime.now(timezone.utc),
            "qa_tested_by": user_id,
        }

        if qa_question is not None:
            update_data["qa_question"] = qa_question

        if qa_answer is not None:
            update_data["qa_answer"] = qa_answer

        if qa_sources is not None:
            update_data["qa_sources_json"] = json.dumps(qa_sources)

        await self._ver_repo.update(version, update_data)

        return doc

    # ── Summary Generation ────────────────────────────────────────────────────

    async def generate_summary_for_document(
        self,
        document_id: UUID,
        current_user,
    ) -> str:
        """
        Generates summary for old documents where summary is missing.

        Flow:
          1. Return existing summary if already available.
          2. If missing, use current version chunks.
          3. Ask AI for detailed document description.
          4. Store summary in document_versions.summary.
          5. Return summary.
        """
        doc = await self._doc_repo.get_with_relations(document_id)

        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")

        current_version = await self._ver_repo.get_current_version(document_id)

        if not current_version:
            raise HTTPException(
                status_code=400,
                detail="Current document version not found",
            )

        if not _is_missing_summary(getattr(current_version, "summary", None)):
            return current_version.summary

        chunks = await self._chunk_repo.get_chunks_by_version(current_version.id)
        chunks = chunks[:12]

        summary_input = "\n\n".join(
            chunk.chunk_text
            for chunk in chunks
            if getattr(chunk, "chunk_text", None)
        )[:12000]

        summary = None

        try:
            summary = await _generate_ai_document_summary(summary_input)

        except Exception as exc:
            logger.warning(
                f"[DocumentService] AI summary generation failed "
                f"for document {document_id}: {exc}"
            )
        if _is_missing_summary(summary):
            summary = (
                getattr(doc, "description", None)
                or summary_input[:1000]
                or "No summary available for this document."
            )

        summary = _clean_markdown_to_plain_text(summary)

        await self._ver_repo.update(
            current_version,
            {
                "summary": summary,
            },
        )

        return summary

    # ── List / Get ────────────────────────────────────────────────────────────

    async def list_documents(
        self,
        page,
        page_size,
        search,
        category_id,
        status_id,
        sort_by,
        sort_order,
    ):
        return await self._doc_repo.list_documents(
            page,
            page_size,
            search,
            category_id,
            status_id,
            sort_by,
            sort_order,
        )
    async def list_published_documents(
        self,
        page,
        page_size,
        search,
        category_id,
        sort_by,
        sort_order,
    ):
        return await self._doc_repo.list_published_documents(
            page,
            page_size,
            search,
            category_id,
            sort_by,
            sort_order,
        )

    async def get_document(self, document_id: UUID) -> Document:
        doc = await self._doc_repo.get_with_relations(document_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found.")
        return doc
    async def list_versions(self, document_id: UUID) -> List[DocumentVersion]:
        await self.get_document(document_id)
        return await self._ver_repo.list_versions(document_id)

    async def get_chunks(self, version_id: UUID) -> List[DocumentChunk]:
        return await self._chunk_repo.get_chunks_by_version(version_id)

    async def get_processing_logs(self, version_id: UUID):
        return await self._log_repo.list_logs_by_version(version_id)

    # ── Update / Delete ───────────────────────────────────────────────────────

    async def update_document(
        self,
        document_id: UUID,
        data: DocumentUpdate,
        user_id: UUID,
    ) -> Document:
        doc = await self.get_document(document_id)

        update_data = {
            key: value
            for key, value in data.model_dump().items()
            if value is not None
        }
        update_data["updated_by"] = user_id

        return await self._doc_repo.update(doc, update_data)

    async def delete_document(self, document_id: UUID) -> Document:
        doc = await self.get_document(document_id)
        return await self._doc_repo.soft_delete(doc)